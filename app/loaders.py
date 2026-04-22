"""Extracción de texto desde archivos subidos (PDF, DOCX, TXT)."""
from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text(file_path: str | Path) -> str:
    """Detecta el tipo de archivo y extrae su texto."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")

    raise ValueError(f"Formato no soportado: {suffix}. Usa PDF, DOCX, TXT o MD.")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            # Marcamos la página para poder citarla luego
            pages.append(f"[Página {i}]\n{text}")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # También incluimos texto de tablas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)
